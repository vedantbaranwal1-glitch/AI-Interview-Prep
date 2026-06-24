from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import text
from sqlalchemy.orm import Session
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from app.gemini import generate_question, evaluate_answer

import os
import shutil
import re
import random

from app.database import SessionLocal, engine
from app.models import Base, User, InterviewHistory, ResumeAnalysis
from app.schemas import UserCreate, UserLogin, InterviewEvaluation
from app.schemas import QuestionRequest

# ==========================================
# Create Database
# ==========================================

Base.metadata.create_all(bind=engine)

# ==========================================
# Upload Folder
# ==========================================

UPLOAD_FOLDER = "uploads"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# ==========================================
# Current Resume Skills
# ==========================================

CURRENT_SKILLS = []
ASKED_QUESTIONS = []

# ==========================================
# FastAPI
# ==========================================

app = FastAPI(
    title="AI Interview Preparation API",
    version="2.0"
)

# ==========================================
# CORS
# ==========================================

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==========================================
# Database Dependency
# ==========================================

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ==========================================
# Root API
# ==========================================

@app.get("/")
def home():
    return {
        "message": "AI Interview Preparation API Running Successfully"
    }

# ==========================================
# Test API
# ==========================================

@app.get("/test")
def test():
    return {
        "message": "API Working Successfully"
    }
@app.get("/dashboard/{user_id}")
def dashboard(user_id: int):

    db = SessionLocal()

    try:

        interviews = db.query(InterviewHistory).filter(
            InterviewHistory.user_id == user_id
        ).all()

        analysis = db.query(ResumeAnalysis).filter(
            ResumeAnalysis.user_id == user_id
        ).first()

        total = len(interviews)

        if total == 0:

            return {
                "total_interviews": 0,
                "average_score": 0,
                "best_score": 0,
                "uploaded_resumes": 1 if analysis else 0,
                "resume_score": analysis.resume_score if analysis else 0,
                "resume_feedback": analysis.resume_feedback if analysis else "No resume analyzed yet."
            }

        scores = [i.score for i in interviews]

        return {
            "total_interviews": total,
            "average_score": round(sum(scores) / total, 2),
            "best_score": max(scores),
            "uploaded_resumes": 1 if analysis else 0,
            "resume_score": analysis.resume_score if analysis else 0,
            "resume_feedback": analysis.resume_feedback if analysis else "No resume analyzed yet."
        }

    finally:
        db.close()
# ==========================================
# Register
# ==========================================

@app.post("/register")
def register(user: UserCreate):

    db = SessionLocal()

    try:

        existing = db.query(User).filter(
            User.email == user.email
        ).first()

        if existing:
            return {
                "success": False,
                "message": "Email already exists"
            }

        new_user = User(
            username=user.username,
            email=user.email,
            password=generate_password_hash(user.password)
        )

        db.add(new_user)
        db.commit()
        db.refresh(new_user)

        return {
            "success": True,
            "message": "Registration Successful",
            "username": new_user.username
        }

    except Exception as e:

        db.rollback()

        return {
            "success": False,
            "message": str(e)
        }

    finally:
        db.close()


# ==========================================
# Login
# ==========================================

@app.post("/login")
def login(user: UserLogin):

    db = SessionLocal()

    try:

        existing = db.query(User).filter(
            User.email == user.email
        ).first()

        if existing is None:
            return {
                "success": False,
                "message": "User not found"
            }

        if not check_password_hash(
            existing.password,
            user.password
        ):
            return {
                "success": False,
                "message": "Wrong password"
            }
        return {
            "success": True,
            "message": "Login Successful",
            "user_id": existing.id,
            "username": existing.username,
            "email": existing.email
        }
    except Exception as e:

        return {
            "success": False,
            "message": str(e)
        }

    finally:
        db.close()
@app.post("/upload-resume")
async def upload_resume(
    user_id: int,
    file: UploadFile = File(...)
):

    print("USER ID:", user_id)

    db = SessionLocal()

    try:

        file_path = os.path.join(
            UPLOAD_FOLDER,
            file.filename
        )

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        result = skill_analysis(file.filename)

        score = result["resume_score"]

        feedback = (
            "Detected Skills: "
            + ", ".join(result["detected_skills"])
            + "\nRecommended Skills: "
            + ", ".join(result["recommended_skills"])
        )

        existing = db.query(ResumeAnalysis).filter(
            ResumeAnalysis.user_id == user_id
        ).first()

        if existing:
            existing.resume_score = score
            existing.resume_feedback = feedback
        else:
            db.add(
                ResumeAnalysis(
                    user_id=user_id,
                    resume_score=score,
                    resume_feedback=feedback
                )
            )

        db.commit()

        print("Resume saved successfully")

        return {
            "success": True,
            "resume_score": score,
            "resume_feedback": feedback
        }

    finally:
        db.close()

    db = SessionLocal()

    try:

        file_path = os.path.join(
            UPLOAD_FOLDER,
            file.filename
        )

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Resume Analyze
        result = skill_analysis(file.filename)

        score = result["resume_score"]

        feedback = (
            "Detected Skills: "
            + ", ".join(result["detected_skills"])
            + "\nRecommended Skills: "
            + ", ".join(result["recommended_skills"])
        )

        # Check if already exists
        existing = db.query(ResumeAnalysis).filter(
            ResumeAnalysis.user_id == user_id
        ).first()

        if existing:
            existing.resume_score = score
            existing.resume_feedback = feedback
        else:
            db.add(
                ResumeAnalysis(
                    user_id=user_id,
                    resume_score=score,
                    resume_feedback=feedback
                )
            )

        db.commit()
        print("Resume saved successfully")

        return {
            "success": True,
            "message": "Resume uploaded successfully",
            "filename": file.filename,
            "resume_score": score,
            "resume_feedback": feedback
        }

    except Exception as e:

        return {
            "success": False,
            "message": str(e)
        }

    finally:
        db.close()
# ==========================================
# Parse Resume
# ==========================================

@app.get("/parse-resume")
def parse_resume(filename: str):

    file_path = os.path.join(
        UPLOAD_FOLDER,
        filename
    )

    if not os.path.exists(file_path):

        return {
            "success": False,
            "message": "Resume not found"
        }

    document = Document(file_path)

    text = ""

    for para in document.paragraphs:
        text += para.text + "\n"

    email = ""

    phone = ""

    email_match = re.search(
        r'[\w\.-]+@[\w\.-]+',
        text
    )

    if email_match:
        email = email_match.group()

    phone_match = re.search(
        r'\+?\d[\d\s-]{8,15}',
        text
    )

    if phone_match:
        phone = phone_match.group()

    skills = []

    skill_list = [

        "Python",
        "Java",
        "C",
        "C++",
        "SQL",
        "React",
        "FastAPI",
        "JavaScript",
        "HTML",
        "CSS",
        "Node",
        "MongoDB",
        "Git",
        "Docker",
        "AWS",
        "Machine Learning",
        "AI",
        "Data Structures",
        "Algorithms"

    ]

    lower_text = text.lower()

    for skill in skill_list:

        if skill.lower() in lower_text:
            skills.append(skill)

    return {

        "success": True,

        "email": email,

        "phone": phone,

        "skills": skills,

        "resume_text": text

    }


# ==========================================
# Skill Analysis
# ==========================================

@app.get("/skill-analysis")
def skill_analysis(filename: str):

    file_path = os.path.join(
        UPLOAD_FOLDER,
        filename
    )

    if not os.path.exists(file_path):

        return {
            "success": False,
            "message": "Resume not found"
        }

    document = Document(file_path)

    text = ""

    for para in document.paragraphs:

        text += para.text + "\n"

    detected = []

    recommended = []

    all_skills = [

        "Python",
        "Java",
        "React",
        "FastAPI",
        "SQL",
        "Docker",
        "Git",
        "AWS",
        "MongoDB",
        "Machine Learning",
        "Communication",
        "Leadership",
        "Problem Solving"

    ]

    for skill in all_skills:

        if skill.lower() in text.lower():

            detected.append(skill)

        else:

            recommended.append(skill)

    score = round(

        len(detected) /

        len(all_skills)

        * 100,

        2

    )
    global CURRENT_SKILLS

    CURRENT_SKILLS = detected
    return {

        "success": True,

        "resume_score": score,

        "detected_skills": detected,

        "recommended_skills": recommended

    }

# ==========================================
# Career Suggestions
# ==========================================

@app.get("/career-suggestions")
def career_suggestions(filename: str):

    file_path = os.path.join(UPLOAD_FOLDER, filename)

    if not os.path.exists(file_path):
        return {
            "success": False,
            "message": "Resume not found"
        }

    document = Document(file_path)

    text = ""

    for para in document.paragraphs:
        text += para.text + "\n"

    text = text.lower()

    careers = []

    if "python" in text:
        careers.append("Python Developer")

    if "react" in text:
        careers.append("Frontend Developer")

    if "fastapi" in text:
        careers.append("Backend Developer")

    if "machine learning" in text or "ai" in text:
        careers.append("AI / Machine Learning Engineer")

    if "sql" in text:
        careers.append("Database Developer")

    if "aws" in text or "docker" in text:
        careers.append("Cloud Engineer")

    if len(careers) == 0:
        careers.append("Software Engineer")

    return {
        "success": True,
        "recommended_careers": careers
    }


@app.post("/evaluate-answer")
def evaluate_answer_api(data: InterviewEvaluation):

    print("STEP 1")

    db = SessionLocal()

    try:

        print("STEP 2")

        result = evaluate_answer(
            data.question,
            data.answer
        )

        print("STEP 3")
        print(result)

        score = result.get("score", 0)
        feedback = result.get("feedback", "")

        interview = InterviewHistory(
            company=data.company,
            role=data.role,
            question=data.question,
            answer=data.answer,
            feedback=feedback,
            score=score
        )

        print("STEP 4")

        db.add(interview)
        db.commit()
        db.refresh(interview)

        print("STEP 5")

        return {
            "success": True,
            "score": score,
            "feedback": feedback
        }

    except Exception as e:
        print("ERROR:", e)
        raise e

    finally:
        db.close()

 
# ==========================================
# Save Interview History
# ==========================================

@app.post("/save-interview")
def save_interview(
    company: str,
    role: str,
    question: str,
    answer: str,
    feedback: str,
    score: int,
):

    db = SessionLocal()

    try:

        interview = InterviewHistory(
            company=company,
            role=role,
            question=question,
            answer=answer,
            feedback=feedback,
            score=score,
        )

        db.add(interview)
        db.commit()
        db.refresh(interview)

        return {
            "success": True,
            "message": "Interview saved successfully"
        }

    except Exception as e:

        db.rollback()

        return {
            "success": False,
            "message": str(e)
        }

    finally:
        db.close()

@app.get("/interview-history/{user_id}")
def interview_history(user_id: int):

    db = SessionLocal()

    try:

        interviews = (
            db.query(InterviewHistory)
            .filter(InterviewHistory.user_id == user_id)
            .order_by(InterviewHistory.id.desc())
            .all()
        )

        history = []

        for item in interviews:

            history.append({
                "id": item.id,
                "company": item.company,
                "role": item.role,
                "question": item.question,
                "answer": item.answer,
                "feedback": item.feedback,
                "score": item.score
            })

        return {
            "success": True,
            "history": history
        }

    finally:
        db.close()
# ==========================================
# Delete Resume
# ==========================================

@app.delete("/delete-resume")
def delete_resume(filename: str):

    file_path = os.path.join(
        UPLOAD_FOLDER,
        filename
    )

    if not os.path.exists(file_path):
        return {
            "success": False,
            "message": "Resume not found"
        }

    os.remove(file_path)

    return {
        "success": True,
        "message": "Resume deleted successfully"
    }


# ==========================================
# Uploaded Resumes
# ==========================================

@app.get("/uploaded-resumes")
def uploaded_resumes():

    resumes = []

    for file in os.listdir(UPLOAD_FOLDER):

        resumes.append(file)

    return {
        "success": True,
        "count": len(resumes),
        "files": resumes
    }


# ==========================================
# Dashboard Statistics
# ==========================================

@app.get("/dashboard-stats")
def dashboard_stats():

    db = SessionLocal()

    try:

        interviews = db.query(
            InterviewHistory
        ).all()

        total = len(interviews)

        if total == 0:

            average = 0
            best = 0

        else:

            scores = [i.score for i in interviews]

            average = round(
                sum(scores) / total,
                2
            )

            best = max(scores)

        resume_count = len(
            os.listdir(UPLOAD_FOLDER)
        )

        return {

            "success": True,

            "total_interviews": total,

            "average_score": average,

            "best_score": best,

            "uploaded_resumes": resume_count

        }

    finally:
        db.close()


# ==========================================
# Health Check
# ==========================================

@app.get("/health")
def health():

    return {

        "status": "Running",

        "database": "Connected",

        "uploads": len(
            os.listdir(UPLOAD_FOLDER)
        )

    }


# ==========================================
# API Information
# ==========================================

@app.get("/about")
def about():

    return {

        "project": "AI Interview Preparation System",

        "version": "2.0",

        "backend": "FastAPI",

        "database": "PostgreSQL",

        "frontend": "React",

        "developer": "Vedant"

    }
@app.post("/evaluate-answer")
def evaluate_answer_api(data: InterviewEvaluation):

    answer = data.answer.lower()

    score = 40

    feedback = []

    if len(answer) > 50:
        score += 15
        feedback.append("Good answer length.")

    if len(answer) > 150:
        score += 10
        feedback.append("Detailed explanation.")

    keywords = [
        "python",
        "java",
        "class",
        "object",
        "inheritance",
        "polymorphism",
        "react",
        "api",
        "database",
        "sql",
        "fastapi",
    ]

    matched = 0

    for word in keywords:
        if word in answer:
            matched += 1

    score += matched * 4

    if score > 100:
        score = 100

    if score >= 90:
        feedback.append("Excellent technical knowledge.")

    elif score >= 75:
        feedback.append("Good answer. Add more examples.")

    else:
        feedback.append("Needs improvement.")

    db = SessionLocal()

    interview = InterviewHistory(

        company=data.company,

        role=data.role,

        question=data.question,

        answer=data.answer,

        feedback=" ".join(feedback),

        score=score,

    )

    db.add(interview)

    db.commit()

    db.refresh(interview)

    db.close()

    return {

        "success": True,

        "score": score,

        "feedback": " ".join(feedback)

    }
@app.post("/generate-question")
def get_question(data: QuestionRequest):
    global CURRENT_SKILLS
    global ASKED_QUESTIONS

    if len(CURRENT_SKILLS) == 0:
        CURRENT_SKILLS = [
            "Python",
            "Communication"
        ]

    if len(ASKED_QUESTIONS) >= 10:
        ASKED_QUESTIONS = []

    selected_skill = random.choice(CURRENT_SKILLS)

    question = generate_question(
        [selected_skill],
        data.company,
        data.difficulty,
        ASKED_QUESTIONS
    )

    ASKED_QUESTIONS.append(question)

    return {
        "success": True,
        "question": question
    }